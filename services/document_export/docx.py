"""글 한 편을 docx 바이트로 굽는다.

`document_create` 도구가 부른다. 본문은 **모델이 쓴 마크다운**인데, 전부를
해석하지 않는다 — 화면 쪽 `AnswerText.tsx` 와 **같은 판단**이다:

    모델이 실제로 쓰는 문법은 몇 개뿐이고(제목 · 목록 · 굵게 · 문단),
    그것만 그리면 의존성이 0이다.

그래서 마크다운 파서를 넣지 않고 줄 단위로 읽는다. 처리하는 것은 넷이다 —
`#`~`###` 제목, `- `/`* ` 목록, `1. ` 번호 목록, `**굵게**`. 나머지는 **글자
그대로** 둔다. 잘못 그리는 것보다 낫고, 그 판단은 화면이 이미 같은 이유로
내렸다.

기존 `body` 문자열은 그대로 지원한다. 다만 보고서처럼 구조가 중요한 문서는
`blocks` 로 제목·문단·목록·안내·표를 명시할 수 있다. 이 구조는 승인 미리보기와
실제 DOCX가 함께 사용하므로, 사용자가 승인한 모양과 생성 파일의 의미가 달라지지
않는다. `table_export` 는 여전히 **표 자체가 결과물인 Excel** 용도이고, 여기의
표는 보고서 안에 들어가는 보조 표다.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

#: `**굵게**` 를 조각으로 가른다. 화면의 `inline()` 이 쓰는 것과 같은 규칙이다.
_BOLD = re.compile(r"(\*\*[^*]+\*\*)")

#: `1. ` · `1) ` 로 시작하는 번호 목록.
_ORDERED = re.compile(r"^\s*\d+[.)]\s+")

#: 제어문자. 워드가 XML 에 못 담는 값이 본문에 섞여 오면 파일이 열리지 않는다.
_CONTROL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")

_NAVY = RGBColor(36, 59, 103)
_BODY = RGBColor(30, 41, 59)
_MUTED = RGBColor(100, 116, 139)
_FONT_LATIN = "Aptos"
_FONT_KOREAN = "맑은 고딕"

SUPPORTED_TEMPLATE_IDS = frozenset({"business_report"})
SUPPORTED_BLOCK_TYPES = frozenset(
    {"heading", "paragraph", "bullet_list", "number_list", "note", "table", "page_break"}
)


def _set_style_font(style, *, size: float, bold: bool = False, color=None) -> None:
    style.font.name = _FONT_LATIN
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    style.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_KOREAN)
    style.element.rPr.rFonts.set(qn("w:ascii"), _FONT_LATIN)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), _FONT_LATIN)
    language = style.element.rPr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        style.element.rPr.append(language)
    language.set(qn("w:val"), "ko-KR")


def _append_page_number(paragraph) -> None:
    paragraph.add_run("페이지 ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend((begin, instruction, separate, result, end))


def _configure_document(document: Document, title: str) -> None:
    """회사 문서로 읽기 위한 최소한의 페이지·서식 체계를 적용한다."""

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    _set_style_font(normal, size=10.5, color=_BODY)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, 14, 6),
        "Heading 2": (13, 12, 5),
        "Heading 3": (11, 10, 4),
    }
    for style_name, (size, before, after) in heading_specs.items():
        style = document.styles[style_name]
        _set_style_font(style, size=size, bold=True, color=_NAVY)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        _set_style_font(style, size=10.5, color=_BODY)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.2

    safe_title = _CONTROL.sub("", title).strip() or "제목 없음"
    document.core_properties.title = safe_title

    header = section.header.paragraphs[0]
    header.text = safe_title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = _FONT_LATIN
        run.font.size = Pt(8)
        run.font.color.rgb = _MUTED
        run._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_KOREAN)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _append_page_number(footer)
    for run in footer.runs:
        run.font.name = _FONT_LATIN
        run.font.size = Pt(8)
        run.font.color.rgb = _MUTED
        run._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_KOREAN)

    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _write_inline(paragraph, text: str) -> None:
    """`**굵게**` 만 살려서 문단에 적는다. 나머지는 글자 그대로."""

    for part in _BOLD.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def _clean(value: Any) -> str:
    return _CONTROL.sub("", "" if value is None else str(value)).strip()


def validate_document_spec(
    *,
    body: str | None = None,
    template_id: str | None = None,
    metadata: dict[str, Any] | None = None,
    blocks: list[dict[str, Any]] | None = None,
) -> None:
    """구조화 문서 입력을 작고 예측 가능한 범위로 제한한다."""

    if template_id and template_id not in SUPPORTED_TEMPLATE_IDS:
        raise ValueError(f"지원하지 않는 문서 템플릿입니다: {template_id}")
    if metadata is not None and not isinstance(metadata, dict):
        raise ValueError("문서 메타데이터는 객체여야 합니다.")
    if blocks is None:
        return
    if not isinstance(blocks, list) or not blocks:
        raise ValueError("구조화 블록은 하나 이상이어야 합니다.")
    if len(blocks) > 100:
        raise ValueError("구조화 블록은 최대 100개까지 만들 수 있습니다.")
    for block in blocks:
        if not isinstance(block, dict) or block.get("type") not in SUPPORTED_BLOCK_TYPES:
            raise ValueError("지원하지 않는 문서 블록이 포함되어 있습니다.")
        if block["type"] in {"heading", "paragraph", "note"} and not _clean(block.get("text")):
            raise ValueError(f"{block['type']} 블록의 text가 비어 있습니다.")
        if block["type"] in {"bullet_list", "number_list"}:
            items = block.get("items")
            if not isinstance(items, list) or not any(_clean(item) for item in items):
                raise ValueError(f"{block['type']} 블록의 items가 비어 있습니다.")
        if block["type"] == "table":
            headers, rows = block.get("headers"), block.get("rows")
            if not isinstance(headers, list) or not headers or len(headers) > 20:
                raise ValueError("문서 표의 열은 1~20개여야 합니다.")
            if not isinstance(rows, list) or len(rows) > 200:
                raise ValueError("문서 표의 행은 최대 200개까지 만들 수 있습니다.")
            if any(not isinstance(row, list) or len(row) > len(headers) for row in rows):
                raise ValueError("문서 표의 각 행은 열 개수를 넘을 수 없습니다.")


def _set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _write_cell(cell, value: Any, *, header: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(_clean(value) or "—")
    run.bold = header
    run.font.name = _FONT_LATIN
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255) if header else _BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_KOREAN)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_metadata(document: Document, metadata: dict[str, Any] | None) -> None:
    labels = (("부서", "department"), ("작성자", "author"), ("작성일", "date"), ("보안 등급", "security"))
    values = [(label, _clean((metadata or {}).get(key))) for label, key in labels]
    values = [(label, value) for label, value in values if value]
    if not values:
        return
    table = document.add_table(rows=len(values), cols=2)
    table.style = "Table Grid"
    for row, (label, value) in zip(table.rows, values, strict=True):
        _set_cell_fill(row.cells[0], "E8EEF8")
        _write_cell(row.cells[0], label)
        row.cells[0].paragraphs[0].runs[0].bold = True
        _write_cell(row.cells[1], value)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_note(document: Document, text: str, label: str | None = None) -> None:
    paragraph = document.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F6FB")
    properties.append(shading)
    if _clean(label):
        label_run = paragraph.add_run(f"{_clean(label)}  ")
        label_run.bold = True
        label_run.font.color.rgb = _NAVY
    _write_inline(paragraph, _clean(text))


def _add_table(document: Document, block: dict[str, Any]) -> None:
    headers = [_clean(value) or "열" for value in block.get("headers", [])]
    rows = block.get("rows", [])
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_fill(cell, "243B67")
        _write_cell(cell, value, header=True)
    for raw_row in rows:
        row = table.add_row()
        # 긴 표가 다음 페이지로 이어지더라도 한 행의 앞뒤가 서로 다른
        # 페이지에 갈라지지 않게 한다. Word가 행 전체를 다음 페이지로
        # 넘기므로 승인 미리보기와 실제 파일의 읽기 흐름이 더 가깝다.
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            _write_cell(cell, raw_row[index] if index < len(raw_row) else None)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _write_blocks(document: Document, blocks: list[dict[str, Any]]) -> None:
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            level = max(1, min(int(block.get("level", 2)), 3))
            document.add_heading(_clean(block.get("text")), level=level)
        elif kind == "paragraph":
            _write_inline(document.add_paragraph(), _clean(block.get("text")))
        elif kind in {"bullet_list", "number_list"}:
            style = "List Bullet" if kind == "bullet_list" else "List Number"
            for item in block.get("items", []):
                if _clean(item):
                    _write_inline(document.add_paragraph(style=style), _clean(item))
        elif kind == "note":
            _add_note(document, _clean(block.get("text")), _clean(block.get("label")))
        elif kind == "table":
            _add_table(document, block)
        elif kind == "page_break":
            document.add_page_break()


def build_docx(
    *,
    title: str,
    body: str = "",
    template_id: str | None = None,
    metadata: dict[str, Any] | None = None,
    blocks: list[dict[str, Any]] | None = None,
) -> bytes:
    """제목과 본문(마크다운)으로 docx 를 만들어 바이트로 돌려준다."""

    validate_document_spec(body=body, template_id=template_id, metadata=metadata, blocks=blocks)
    document = Document()

    _configure_document(document, title)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(_CONTROL.sub("", title).strip() or "제목 없음")
    run.bold = True
    run.font.name = _FONT_LATIN
    run.font.size = Pt(22)
    run.font.color.rgb = _NAVY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_KOREAN)
    heading.paragraph_format.space_after = Pt(16)
    heading.paragraph_format.keep_with_next = True

    if template_id == "business_report":
        _add_metadata(document, metadata)

    if blocks is not None:
        _write_blocks(document, blocks)
    else:
        for raw in _CONTROL.sub("", body or "").split("\n"):
            line = raw.rstrip()
            stripped = line.strip()

            # 빈 줄은 문단을 나눈다. 빈 문단을 넣지는 않는다 — 워드에서 여백이
            # 두 배로 벌어진다(문단 간격이 이미 있다).
            if not stripped:
                continue

            if stripped.startswith("### "):
                document.add_heading(stripped[4:].strip(), level=3)
            elif stripped.startswith("## "):
                document.add_heading(stripped[3:].strip(), level=2)
            elif stripped.startswith("# "):
                document.add_heading(stripped[2:].strip(), level=1)
            elif stripped.startswith(("- ", "* ")):
                _write_inline(document.add_paragraph(style="List Bullet"), stripped[2:].strip())
            elif _ORDERED.match(stripped):
                _write_inline(
                    document.add_paragraph(style="List Number"), _ORDERED.sub("", stripped)
                )
            else:
                _write_inline(document.add_paragraph(), stripped)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

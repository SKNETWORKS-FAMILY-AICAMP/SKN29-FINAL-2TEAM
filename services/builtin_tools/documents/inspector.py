"""파일 바이트의 기본 정보와 포맷별 구조를 확인한다."""

from __future__ import annotations

import hashlib
from io import BytesIO
from typing import Any
from zipfile import BadZipFile

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from services.builtin_tools.common.errors import BuiltinToolError
from services.builtin_tools.common.limits import MAX_FILE_BYTES
from services.builtin_tools.documents.reader import DOCX_MIME, PDF_MIME, XLSX_MIME
from services.builtin_tools.documents.package_safety import validate_ooxml_package


def detect_mime_type(data: bytes) -> str:
    """확장자나 사용자 선언값이 아니라 파일 signature로 MIME을 판별한다."""

    try:
        # python-magic은 import 시점에 OS의 libmagic을 즉시 로드한다. Windows처럼
        # 네이티브 라이브러리가 없는 개발 환경에서도 Django 자체는 시작돼야 하므로
        # 실제 MIME 판별을 요청할 때까지 로드를 미룬다.
        import magic

        detected = str(magic.from_buffer(data, mime=True))
    except (ImportError, OSError) as exc:
        raise BuiltinToolError(
            "MIME_DETECTOR_UNAVAILABLE",
            "파일 형식 판별기를 사용할 수 없습니다. libmagic 설치 상태를 확인해 주세요.",
        ) from exc
    except Exception as exc:
        # python-magic의 예외 타입은 모듈 로드 후에만 참조할 수 있다. 라이브러리
        # 내부 오류를 외부로 노출하지 않고 기존 도구 오류 계약으로 변환한다.
        if exc.__class__.__module__.startswith("magic"):
            raise BuiltinToolError(
                "MIME_DETECTION_FAILED", "파일 형식을 확인하지 못했습니다."
            ) from exc
        raise
    return detected.split(";", 1)[0].strip().lower()


def _metadata(properties: Any, fields: tuple[str, ...]) -> dict[str, Any]:
    result = {}
    for field in fields:
        value = getattr(properties, field, None)
        if value not in (None, ""):
            result[field] = value.isoformat() if hasattr(value, "isoformat") else str(value)
    return result


def _inspect_pdf(data: bytes) -> dict[str, Any]:
    try:
        reader = PdfReader(BytesIO(data))
    except PdfReadError as exc:
        raise BuiltinToolError("INVALID_PDF", "손상되었거나 지원하지 않는 PDF입니다.") from exc
    if reader.is_encrypted:
        return {"page_count": None, "encrypted": True, "metadata": {}}
    return {
        "page_count": len(reader.pages),
        "encrypted": False,
        "metadata": {
            str(key).lstrip("/"): str(value) for key, value in (reader.metadata or {}).items()
        },
    }


def _inspect_docx(data: bytes) -> dict[str, Any]:
    validate_ooxml_package(data)
    try:
        document = Document(BytesIO(data))
    except (ValueError, KeyError, BadZipFile, InvalidFileException) as exc:
        raise BuiltinToolError("INVALID_DOCX", "손상되었거나 지원하지 않는 DOCX입니다.") from exc
    return {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "metadata": _metadata(
            document.core_properties,
            (
                "title",
                "subject",
                "author",
                "last_modified_by",
                "keywords",
                "category",
                "comments",
                "created",
                "modified",
            ),
        ),
    }


def _inspect_xlsx(data: bytes) -> dict[str, Any]:
    validate_ooxml_package(data)
    try:
        workbook = load_workbook(BytesIO(data), read_only=True, data_only=False)
    except (ValueError, KeyError, BadZipFile) as exc:
        raise BuiltinToolError("INVALID_XLSX", "손상되었거나 지원하지 않는 XLSX입니다.") from exc
    try:
        return {
            "sheet_count": len(workbook.sheetnames),
            "sheets": [
                {"name": sheet.title, "rows": sheet.max_row, "columns": sheet.max_column}
                for sheet in workbook.worksheets
            ],
            "metadata": _metadata(
                workbook.properties,
                (
                    "title",
                    "subject",
                    "creator",
                    "lastModifiedBy",
                    "keywords",
                    "category",
                    "description",
                    "created",
                    "modified",
                ),
            ),
        }
    finally:
        workbook.close()


def inspect_file(*, data: bytes, mime_type: str) -> dict[str, Any]:
    """실제 MIME, hash, 크기와 포맷별 속성을 반환한다."""

    if not data:
        raise BuiltinToolError("EMPTY_FILE", "빈 파일은 검사할 수 없습니다.")
    if len(data) > MAX_FILE_BYTES:
        raise BuiltinToolError("FILE_TOO_LARGE", "파일은 20MB 이하여야 합니다.")
    detected_mime_type = detect_mime_type(data)
    if detected_mime_type == PDF_MIME:
        details = _inspect_pdf(data)
    elif detected_mime_type == DOCX_MIME:
        details = _inspect_docx(data)
    elif detected_mime_type == XLSX_MIME:
        details = _inspect_xlsx(data)
    else:
        details = {}
    return {
        "size_bytes": len(data),
        "sha256": hashlib.sha256(data).hexdigest(),
        "declared_mime_type": mime_type,
        "detected_mime_type": detected_mime_type,
        "mime_verified": detected_mime_type == mime_type,
        **details,
    }

"""PDF·DOCX·XLSX의 내용을 정해진 구조로 읽는다."""

from __future__ import annotations

from io import BytesIO
from typing import Any
from zipfile import BadZipFile

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException
import pdfplumber
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from services.builtin_tools.common.errors import BuiltinToolError
from services.builtin_tools.common.json_values import json_value
from services.builtin_tools.common.isolation import run_parser_isolated
from services.builtin_tools.common.limits import (
    MAX_FILE_BYTES,
    MAX_PDF_PAGES,
    MAX_WORKBOOK_CELLS,
    MAX_WORKBOOK_SHEETS,
)
from services.builtin_tools.documents.package_safety import validate_ooxml_package

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
SUPPORTED_MIME_TYPES = frozenset({PDF_MIME, DOCX_MIME, XLSX_MIME})


def _check_input(data: bytes, mime_type: str) -> None:
    if not data:
        raise BuiltinToolError("EMPTY_FILE", "빈 파일은 읽을 수 없습니다.")
    if len(data) > MAX_FILE_BYTES:
        raise BuiltinToolError("FILE_TOO_LARGE", "파일은 20MB 이하여야 합니다.")
    if mime_type not in SUPPORTED_MIME_TYPES:
        raise BuiltinToolError("UNSUPPORTED_FORMAT", "PDF, DOCX, XLSX 파일만 읽을 수 있습니다.")


def _read_pdf_in_process(*, data: bytes) -> dict[str, Any]:
    try:
        reader = PdfReader(BytesIO(data))
    except PdfReadError as exc:
        raise BuiltinToolError("INVALID_PDF", "손상되었거나 지원하지 않는 PDF입니다.") from exc
    if reader.is_encrypted:
        raise BuiltinToolError("PASSWORD_REQUIRED", "암호가 설정된 PDF는 읽을 수 없습니다.")
    if len(reader.pages) > MAX_PDF_PAGES:
        raise BuiltinToolError("PAGE_LIMIT_EXCEEDED", "PDF는 200페이지까지 읽을 수 있습니다.")

    sections = []
    tables = []
    warnings = []
    try:
        with pdfplumber.open(BytesIO(data)) as document:
            for number, page in enumerate(document.pages, start=1):
                try:
                    text = (page.extract_text() or "").strip()
                except Exception:  # noqa: BLE001 - 원문을 숨기고 페이지 경고만 반환한다.
                    text = ""
                    warnings.append({"code": "TEXT_EXTRACTION_FAILED", "page": number})
                # pdfplumber가 글자를 못 읽었을 때만 이미 검증된 pypdf로 한 번 더 본다.
                if not text:
                    try:
                        text = (reader.pages[number - 1].extract_text() or "").strip()
                    except Exception:  # noqa: BLE001 - 두 parser 모두 실패하면 경고로 끝낸다.
                        text = ""
                sections.append({"location": {"page": number}, "text": text})
                if not text and not any(item["page"] == number for item in warnings):
                    warnings.append({"code": "TEXT_NOT_EXTRACTABLE", "page": number})
                try:
                    extracted = page.extract_tables() or []
                except Exception:  # noqa: BLE001 - 본문은 살리고 표만 신뢰 불가로 표시한다.
                    extracted = []
                    warnings.append({"code": "TABLE_EXTRACTION_FAILED", "page": number})
                for table_number, rows in enumerate(extracted, start=1):
                    tables.append(
                        {
                            "location": {"page": number, "table": table_number},
                            "rows": rows,
                        }
                    )
    except Exception as exc:  # noqa: BLE001 - pdfplumber 내부 원문·경로를 노출하지 않는다.
        raise BuiltinToolError("INVALID_PDF", "PDF 내용을 읽지 못했습니다.") from exc
    return {"format": "pdf", "sections": sections, "tables": tables, "warnings": warnings}


def _read_pdf(data: bytes) -> dict[str, Any]:
    return run_parser_isolated(
        module_name=__name__, function_name="_read_pdf_in_process", kwargs={"data": data}
    )


def _read_docx(data: bytes) -> dict[str, Any]:
    validate_ooxml_package(data)
    try:
        document = Document(BytesIO(data))
    except (ValueError, KeyError, BadZipFile) as exc:
        raise BuiltinToolError("INVALID_DOCX", "손상되었거나 지원하지 않는 DOCX입니다.") from exc

    sections = [
        {"location": {"paragraph": index}, "text": paragraph.text}
        for index, paragraph in enumerate(document.paragraphs, start=1)
        if paragraph.text.strip()
    ]
    tables = []
    for table_number, table in enumerate(document.tables, start=1):
        tables.append(
            {
                "location": {"table": table_number},
                "rows": [[cell.text for cell in row.cells] for row in table.rows],
            }
        )
    return {"format": "docx", "sections": sections, "tables": tables, "warnings": []}


def _read_xlsx(data: bytes) -> dict[str, Any]:
    validate_ooxml_package(data)
    try:
        workbook = load_workbook(BytesIO(data), read_only=True, data_only=False)
    except (ValueError, KeyError, BadZipFile, InvalidFileException) as exc:
        raise BuiltinToolError("INVALID_XLSX", "손상되었거나 지원하지 않는 XLSX입니다.") from exc
    try:
        if len(workbook.sheetnames) > MAX_WORKBOOK_SHEETS:
            raise BuiltinToolError("SHEET_LIMIT_EXCEEDED", "Excel은 100개 시트까지 읽을 수 있습니다.")

        tables = []
        cell_count = 0
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows():
                cell_count += len(row)
                if cell_count > MAX_WORKBOOK_CELLS:
                    raise BuiltinToolError(
                        "CELL_LIMIT_EXCEEDED", "Excel은 전체 500,000개 셀까지 읽을 수 있습니다."
                    )
                rows.append([json_value(cell.value) for cell in row])
            tables.append({"location": {"sheet": sheet.title}, "rows": rows})
        return {"format": "xlsx", "sections": [], "tables": tables, "warnings": []}
    finally:
        workbook.close()


def read_document(*, data: bytes, mime_type: str) -> dict[str, Any]:
    """파일 바이트를 읽고 형식·본문·표·경고를 반환한다."""

    _check_input(data, mime_type)
    if mime_type == PDF_MIME:
        result = _read_pdf(data)
    elif mime_type == DOCX_MIME:
        result = _read_docx(data)
    else:
        result = _read_xlsx(data)
    return {**result, "truncated": False}

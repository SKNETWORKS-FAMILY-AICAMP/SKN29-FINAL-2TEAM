"""DOCX·XLSX의 구조를 보존하는 파일 비교."""

from __future__ import annotations

from difflib import SequenceMatcher
from io import BytesIO
from typing import Any
from zipfile import BadZipFile

from docx import Document
from openpyxl import load_workbook
from services.builtin_tools.common.errors import BuiltinToolError
from services.builtin_tools.common.json_values import json_value
from services.builtin_tools.common.limits import MAX_FILE_BYTES
from services.builtin_tools.documents.reader import (
    DOCX_MIME,
    PDF_MIME,
    XLSX_MIME,
    read_document,
)
from services.builtin_tools.documents.package_safety import validate_ooxml_package


def _sequence_changes(before: list[str], after: list[str], *, location: str) -> list[dict[str, Any]]:
    changes = []
    matcher = SequenceMatcher(a=before, b=after, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        changes.append(
            {
                "type": {"insert": "added", "delete": "removed", "replace": "modified"}[tag],
                "location": location,
                "before_range": [i1 + 1, i2],
                "after_range": [j1 + 1, j2],
                "before": before[i1:i2],
                "after": after[j1:j2],
            }
        )
    return changes


def _docx_lines(data: bytes) -> tuple[list[str], list[str]]:
    validate_ooxml_package(data)
    try:
        document = Document(BytesIO(data))
    except (ValueError, KeyError, BadZipFile) as exc:
        raise BuiltinToolError("INVALID_DOCX", "손상되었거나 지원하지 않는 DOCX입니다.") from exc
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        f"표{table_index}/행{row_index}/열{cell_index}: {cell.text}"
        for table_index, table in enumerate(document.tables, start=1)
        for row_index, row in enumerate(table.rows, start=1)
        for cell_index, cell in enumerate(row.cells, start=1)
    ]
    return paragraphs, table_cells


def _xlsx_cells(data: bytes) -> tuple[set[str], dict[str, Any]]:
    validate_ooxml_package(data)
    try:
        workbook = load_workbook(BytesIO(data), read_only=True, data_only=False)
    except (ValueError, KeyError, BadZipFile) as exc:
        raise BuiltinToolError("INVALID_XLSX", "손상되었거나 지원하지 않는 XLSX입니다.") from exc
    try:
        cells = {}
        sheet_names = set(workbook.sheetnames)
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        cells[f"{sheet.title}!{cell.coordinate}"] = json_value(cell.value)
        return sheet_names, cells
    finally:
        workbook.close()


def _compare_xlsx(before: bytes, after: bytes) -> list[dict[str, Any]]:
    old_sheets, old = _xlsx_cells(before)
    new_sheets, new = _xlsx_cells(after)
    changes = [
        {"type": "sheet_removed", "location": sheet, "before": sheet, "after": None}
        for sheet in sorted(old_sheets - new_sheets)
    ]
    changes += [
        {"type": "sheet_added", "location": sheet, "before": None, "after": sheet}
        for sheet in sorted(new_sheets - old_sheets)
    ]
    for location in sorted(old.keys() | new.keys()):
        if location not in old:
            changes.append({"type": "added", "location": location, "before": None, "after": new[location]})
        elif location not in new:
            changes.append({"type": "removed", "location": location, "before": old[location], "after": None})
        elif old[location] != new[location]:
            changes.append({"type": "modified", "location": location, "before": old[location], "after": new[location]})
    return changes


def _pdf_pages(data: bytes) -> list[str]:
    result = read_document(data=data, mime_type=PDF_MIME)
    return [section["text"] for section in result["sections"]]


def compare_files(*, before: bytes, after: bytes, mime_type: str) -> dict[str, Any]:
    """같은 형식의 두 파일 변경점을 반환한다."""

    if not before or not after:
        raise BuiltinToolError("EMPTY_FILE", "비교할 두 파일이 모두 필요합니다.")
    if len(before) > MAX_FILE_BYTES or len(after) > MAX_FILE_BYTES:
        raise BuiltinToolError("FILE_TOO_LARGE", "비교할 파일은 각각 20MB 이하여야 합니다.")
    if mime_type == DOCX_MIME:
        before_paragraphs, before_tables = _docx_lines(before)
        after_paragraphs, after_tables = _docx_lines(after)
        changes = _sequence_changes(before_paragraphs, after_paragraphs, location="paragraphs")
        changes += _sequence_changes(before_tables, after_tables, location="tables")
    elif mime_type == XLSX_MIME:
        changes = _compare_xlsx(before, after)
    elif mime_type == PDF_MIME:
        changes = _sequence_changes(_pdf_pages(before), _pdf_pages(after), location="pages")
    else:
        raise BuiltinToolError(
            "UNSUPPORTED_FORMAT", "PDF, DOCX, XLSX 파일만 비교할 수 있습니다."
        )
    output_format = {DOCX_MIME: "docx", XLSX_MIME: "xlsx", PDF_MIME: "pdf"}[mime_type]
    return {"format": output_format, "change_count": len(changes), "changes": changes}

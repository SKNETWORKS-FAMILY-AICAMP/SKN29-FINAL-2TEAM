from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"D:\Project\SKN29-Final-2Team\최종발표_PPT_기획정리.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_BLUE = "F4F7FA"
GRAY = "666666"
FONT = "Malgun Gothic"


def set_font(run, size=None, bold=None, color=None, name=FONT):
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_text(paragraph, text, bold=False, color=None, size=11):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.25
    pf.left_indent = Inches(0.375 + (0.25 * level))
    pf.first_line_indent = Inches(-0.188)
    add_text(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.25
    pf.left_indent = Inches(0.375)
    pf.first_line_indent = Inches(-0.188)
    add_text(p, text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_text(p, text, bold=True, color=BLUE if level < 3 else DARK_BLUE,
             size={1: 16, 2: 13, 3: 12}[level])
    set_keep_with_next(p)
    return p


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    mark_header_row(table.rows[0])
    shade_cell(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, f"{label}  ", bold=True, color=DARK_BLUE)
    add_text(p, body)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for cell, value in zip(table.rows[0].cells, headers):
        shade_cell(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, value, bold=True, color=DARK_BLUE, size=10)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, value, size=9.5)
    return table


def add_section_break(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    add_text(p, "최종 발표 PPT 기획 정리 | 2026.08.30 멘토링", color=GRAY, size=8.5)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_footer(doc)

    # Opening block - compact reference guide, named Korean-font override.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_text(p, "MENTORING OUTPUT", bold=True, color=BLUE, size=10)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_text(p, "최종 발표 PPT 기획 정리", bold=True, color=DARK_BLUE, size=24)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    add_text(p, "2026.08.30 온라인 멘토링 회의록 기반 | 발표 총괄 및 파트별 콘텐츠 전달용", color=GRAY, size=10)

    add_callout(doc, "문서 목적", "회의록 전체에서 최종 발표 PPT의 메시지, 장표 구성, 증거 자료, 수정 지시, 한계 및 준비 액션만 분리해 바로 장표화할 수 있도록 정리한다.")

    add_heading(doc, "1. 발표의 한 문장", 1)
    p = doc.add_paragraph()
    add_text(p, "우리는 ", size=12)
    add_text(p, "비개발자도 Agent를 구성·운영할 수 있는 플랫폼", bold=True, color=DARK_BLUE, size=12)
    add_text(p, "에 문서 파싱 보강과 실행 검증 체계를 결합해, 실제 업무 문서 기반의 신뢰 가능한 자동화를 제공한다.", size=12)

    add_heading(doc, "2. 권장 발표 스토리라인", 1)
    storyline = [
        ("1", "시장과 포지셔닝", "왜 Agent Platform이 필요한지, 국내·글로벌 경쟁 제품과 비교해 우리 서비스의 빈틈을 제시"),
        ("2", "문제와 핵심 경험", "비개발자가 Agent를 정의하고 Tool을 연결해 실제 Chat 실행까지 이어지는 경험을 짧게 소개"),
        ("3", "Deep Agent 구조", "Agent Builder, Tool 조합, Versioning, 실행 구조를 한 아키텍처로 설명"),
        ("4", "문서 파싱 차별화", "기본 Docling의 실패 사례를 Gate별 보완 파이프라인과 Before/After로 제시"),
        ("5", "이미지·멀티모달 연결", "Description 검색에서 Metadata·원본 Crop·VLM 답변으로 이어지는 설계 의도와 운영 조건 제시"),
        ("6", "플랫폼 평가", "모델 성능 경쟁이 아니라 Tool 호출·종료·안전성·운영 효율을 검증한 Evidence 제시"),
        ("7", "운영과 확장", "Admin 실제 화면, Skill의 현재 범위, Versioning의 Production 확장성을 제시"),
        ("8", "한계와 다음 단계", "현재 구현 범위를 정직하게 밝히고 Future Work 및 QA·Demo 안정화 계획으로 마무리"),
    ]
    add_matrix(doc, ["순서", "장표 주제", "핵심 전달 내용"], storyline, [700, 1900, 6760])

    add_heading(doc, "3. 장표별 메시지와 보여줄 근거", 1)
    sections = [
        ("3-1. 시장 조사·포지셔닝", [
            "프로젝트의 출발점은 기능 나열이 아니라, Agent Platform/업무 자동화 시장에서 사용자가 겪는 구성·운영·문서 활용의 단절을 해결하는 데 있다.",
            "국내·글로벌 유사 제품을 조사해 각 제품의 강점과 비어 있는 지점을 한 장표로 연결한다.",
            "단순 기능 비교표가 아니라 ‘우리 서비스가 들어갈 자리’를 보여주는 결론 중심 장표로 구성한다.",
        ], "경쟁 제품 포지셔닝 맵 또는 간결한 비교표"),
        ("3-2. Agent Builder와 Deep Agent Architecture", [
            "상태 3개, 도구 32개 등 맥락 없는 숫자 강조는 제거한다.",
            "비개발자가 Agent를 정의하고, 필요한 Tool을 연결하고, Version을 관리한 뒤 실제 Chat에서 실행하는 흐름을 핵심으로 둔다.",
            "Versioning은 Dev/Stage/Prod 수준의 운영 확장성을 설명할 수 있으므로 유지한다.",
            "Agent 영역에는 Deep Agent Architecture를 직접 넣어 Builder·Tool·실행의 관계를 한눈에 보여준다.",
        ], "Builder 흐름 화면 + Versioning 화면 + Deep Agent Architecture"),
        ("3-3. 문서 연결과 인덱싱 UX", [
            "커넥터 연결 후 문서가 수집·인덱싱되고, 사용자가 상태를 확인할 수 있음을 보여준다.",
            "문서 페이지를 노출한다면 인덱싱 상태 확인을 차별점으로 두고, 갱신/재동기화 방향은 후속 UX로 연결한다.",
        ], "문서 페이지 실제 Screenshot + 인덱싱 상태"),
        ("3-4. 문서 파싱: Gate별 문제 → 보완 → 결과 → 평가", [
            "단순 Architecture 소개가 아니라 기본 Docling 결과의 실패 사례와 각 보완 Layer의 역할을 대응시킨다.",
            "Reading Order/Heading, Table Gate, Picture 보강, Description Quality Gate, Chunking·Serialization의 순서로 설명한다.",
            "대표 파라미터는 1~2개만 설명한다. 예: max_new_tokens=224로 반복·Context Echo 완화, temperature=0·do_sample=false로 일관성 확보.",
            "이 값은 제한된 브로슈어 샘플을 기반으로 한 휴리스틱 실험값이며 대규모 데이터에서 재튜닝이 필요하다고 먼저 밝힌다.",
        ], "Before/After 예시 + Pipeline 그림 + Docling 대비 처리시간·품질 비교"),
        ("3-5. 이미지 검색과 멀티모달 답변", [
            "Picture Description은 검색용 텍스트로 임베딩하고, 검색된 후보는 문서·페이지·BBox 메타데이터로 원본 Crop을 역추적한다.",
            "최종 답변에는 Description만이 아니라 Crop Image를 Vision/VLM 모델에 전달하는 것이 설계 목표다.",
            "모든 OpenAI-compatible 모델이 이미지 입력을 지원하지 않으므로, Admin/Runtime에 Vision Capability Check가 필요하다.",
        ], "Description → Metadata → Crop → VLM 흐름도; Capability 표시 화면"),
        ("3-6. 플랫폼 평가", [
            "평가 대상은 특정 LLM의 지능이 아니라 사용자가 만든 Agent가 플랫폼에서 올바르고 안전하며 효율적으로 실행되는지다.",
            "정상 종료, 답변 존재, 필수·허용 Tool 사용, 중복 호출, 문서 사용·데이터 유출 조건을 결정론적 Rule/Assertion으로 확인한다.",
            "동일 시나리오를 3회 반복해 변동을 기록하고, LLM-as-a-Judge·Phoenix/OpenTelemetry·garak은 각각 보조 품질·관측·보안 검증 역할로 구분한다.",
            "발표에서는 DB 컬럼과 영문 Evaluator 대신 쉬운 한국어로 ‘평가 목적 → 기준 → 결과 → 의미’를 보여준다.",
        ], "Rule 기반 항목 + 반복 시나리오 결과 + Phoenix Trace 요약"),
        ("3-7. Skill과 Admin 운영", [
            "현재 Skill은 생성 → Store 저장 → 팀 공유 → / 명시 호출이 가능한 Skill/Shortcut이다.",
            "Query 의도에 따른 자동 Trigger/Hook은 아직 구현 범위가 아니므로 Future Work로 분리한다.",
            "Admin은 실행·모델·Tool 상태를 관리하는 실제 화면을 제시한다. 근거 없는 최근 30일·실행 수·Tool 수는 제거한다.",
            "실제 테스트 규모를 제시할 때는 개발 과정의 증거라는 맥락과 함께 쓴다.",
        ], "Skill 생성/공유 화면 + 현재 한계/Future Work + Admin 실제 Screenshot"),
    ]
    for title, bullets, evidence in sections:
        add_heading(doc, title, 2)
        for item in bullets:
            add_bullet(doc, item)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        add_text(p, "권장 Evidence: ", bold=True, color=DARK_BLUE, size=10)
        add_text(p, evidence, size=10)

    add_section_break(doc)
    add_heading(doc, "4. 반드시 반영할 수정·삭제 기준", 1)
    add_matrix(doc, ["구분", "반영 지침"], [
        ("추가", "시장 조사·경쟁 제품 포지셔닝, Deep Agent Architecture, Admin 실제 Screenshot"),
        ("압축", "초반 2번 장표는 약 30초 안에 설명하고 넘어갈 수준으로 간결화"),
        ("유지", "Versioning의 운영 확장성, Preamble 기반 작업 과정 표시"),
        ("표현 수정", "Skill은 ‘자동 Trigger Agent Skill’이 아니라 ‘명시적 호출형 Skill/Shortcut’으로 설명"),
        ("삭제", "상태 3개·도구 32개, Draft/Active/Disabled 세부 상태, 근거 없는 최근 30일·실행 수·Tool 수"),
        ("주의", "Candidate Model 이름이나 모델 점수를 중심에 두지 않는다. 모델 성능과 플랫폼 동작 평가는 분리한다."),
    ], [1900, 7460])

    add_heading(doc, "5. 발표에서 먼저 인정할 한계와 Future Work", 1)
    future = [
        "문장 단위 Citation: 답변의 주장과 근거 문서를 직접 연결",
        "Table False Negative: 도형으로 그린 표 등 실제 표를 인식하지 못하는 사례 보완",
        "Description Gate 과삭제: 정상 이미지 설명의 과삭제를 줄일 품질 기준과 Fallback 마련",
        "모델 Capability Routing: Vision 지원 모델에만 Crop Image를 전달하도록 Admin/Runtime 연계",
        "Table Serialization Routing: Markdown·Triplet·Crop+VLM을 표 복잡도에 따라 선택",
        "MCP Context Budget: Unique Tool 수, Loop, Token/Context 사용량을 제어·요약",
        "자동 Skill Trigger: 명시 호출 외 Query 의도 기반 Hook/Router 구현",
        "대규모 Evaluation: 제한된 샘플의 휴리스틱을 다양한 문서·질문으로 검증",
    ]
    for item in future:
        add_bullet(doc, item)

    add_heading(doc, "6. PPT 제작 전 콘텐츠 취합 액션", 1)
    add_matrix(doc, ["담당", "발표용 전달물", "목표"], [
        ("지훈 + 주연", "Deep Agent 발표 Storyline", "8/31 멘토링 전 정리"),
        ("지훈", "평가 지표·결과의 쉬운 한국어 요약", "8/31 멘토링 전"),
        ("원빈", "문서 처리의 핵심 내용·Evidence·Failure Case", "8/31 멘토링 전"),
        ("PPT 총괄", "담당자 자료를 하나의 스토리라인으로 통합", "자료 취합 후"),
        ("전체 팀", "QA·회귀 테스트·Demo 안정화", "지속"),
        ("원빈 중심 + 팀", "최종 시나리오 시연 영상", "자료 정리와 QA 이후"),
    ], [1700, 5100, 2560])

    add_callout(doc, "운영 원칙", "담당자별로 완성된 디자인 슬라이드를 따로 만드는 방식이 아니라, 각 파트는 근거와 Storyline을 Markdown 등 바로 옮길 수 있는 형태로 전달하고 PPT 총괄이 하나의 흐름으로 통합한다.")

    add_heading(doc, "7. 최종 점검 체크리스트", 1)
    checklist = [
        "각 장표가 ‘무엇을 보여주는가’보다 ‘왜 이 기능이 중요한가’라는 메시지에 답하는가",
        "모든 숫자와 지표에 출처·기간·측정 조건 또는 개발 과정의 맥락이 있는가",
        "문서 파싱은 Gate별 Failure Case와 Before/After가 연결되어 있는가",
        "평가는 모델 점수 대신 플랫폼 동작 검증이라는 범위를 분명히 했는가",
        "현재 구현, 설계 목표, Future Work를 혼동하지 않았는가",
        "Demo와 Screenshot은 실제 동작·현재 화면과 일치하며 QA가 완료되었는가",
    ]
    for item in checklist:
        add_bullet(doc, item)

    doc.core_properties.title = "최종 발표 PPT 기획 정리"
    doc.core_properties.subject = "2026-08-30 멘토링 회의록 기반 최종 발표 PPT 관련 내용"
    doc.core_properties.author = "SKN29 Final 2Team"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

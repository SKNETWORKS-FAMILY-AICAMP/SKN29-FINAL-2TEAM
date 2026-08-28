"""신규 기본 도구 10종의 localhost 실측용 「내 파일」 픽스처를 만든다.

`web` 컨테이너에서 실행한다:

    docker exec skn29-final-2team-web-1 python scripts/seed_builtin_tool_fixtures.py [account_email]

- 기본 대상은 `SEED_ACCOUNT_EMAIL` 환경변수 또는 인자로 준 이메일, 없으면
  가장 먼저 만들어진 ACTIVE 계정.
- 같은 이름의 기존 픽스처(source_type='UPLOAD')는 지우고 다시 만든다 — 여러 번
  돌려도 목록이 불어나지 않는다.
- 실제 서비스 코드(`PersonalDocumentRepository`, `storage`,
  `office_to_pdf`)를 그대로 써서 도구가 받는 것과 같은 파일을 만든다.
"""

from __future__ import annotations

from io import BytesIO
import json
import os
from pathlib import Path
import sys
import zipfile

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")

import django  # noqa: E402

django.setup()

from docx import Document  # noqa: E402
from openpyxl import Workbook  # noqa: E402

from backend.db.connection import database_connection  # noqa: E402
from backend.db.document_pipeline import PersonalDocumentRepository  # noqa: E402
from backend.db.repositories import DocumentRepository  # noqa: E402
from backend.services import storage  # noqa: E402
from services.builtin_tools.documents.converter import office_to_pdf  # noqa: E402

_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_PDF = "application/pdf"
_CSV = "text/csv"
_JSON = "application/json"
_ZIP = "application/zip"
_MD = "text/markdown"
_PARQUET = "application/vnd.apache.parquet"


def _report_docx(*, author: str, revised: bool = False) -> bytes:
    document = Document()
    document.core_properties.author = author
    document.core_properties.last_modified_by = author
    document.core_properties.title = "2026년 2분기 사업 보고서"
    document.core_properties.subject = "분기 실적"
    document.core_properties.keywords = "실적, 매출, 분기"
    document.add_heading("2026년 2분기 사업 보고서", level=1)
    document.add_paragraph("작성 부서: 전략기획팀")
    if revised:
        document.add_paragraph(
            "2분기 매출은 전분기 대비 18퍼센트 증가했으며, 신규 계약 6건이 체결되었다."
        )
        document.add_paragraph("3분기에는 해외 파트너십 2건을 추가로 추진한다.")
    else:
        document.add_paragraph(
            "2분기 매출은 전분기 대비 12퍼센트 증가했으며, 신규 계약 4건이 체결되었다."
        )
    document.add_heading("부서별 요약", level=2)
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    dev_result = "릴리스 3회" if revised else "릴리스 2회"
    rows = [("부서", "핵심 성과"), ("영업", "신규 계약 4건"), ("개발", dev_result)]
    for r, (a, b) in enumerate(rows):
        table.rows[r].cells[0].text = a
        table.rows[r].cells[1].text = b
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _headcount_xlsx() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "부서별공수"
    sheet.append(["부서", "이름", "공수"])
    for row in [
        ["영업", "김영업", 40],
        ["영업", "이세일", 32],
        ["개발", "박개발", 45],
        ["개발", "최코드", 38],
        ["개발", "정빌드", 41],
        ["디자인", "한디자", 30],
    ]:
        sheet.append(row)
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _roster_csv() -> bytes:
    # 빈 값 1건(3행 이메일), 완전 중복 1건(김영업 두 번).
    lines = [
        "사번,이름,부서,이메일",
        "E001,김영업,영업,kim@example.com",
        "E002,이세일,영업,",
        "E003,박개발,개발,park@example.com",
        "E001,김영업,영업,kim@example.com",
        "E004,최코드,개발,choi@example.com",
    ]
    return ("\n".join(lines) + "\n").encode("utf-8")


def _config_json() -> bytes:
    # retries 가 문자열이라 "정수" 스키마를 주면 위반이 잡힌다.
    return json.dumps(
        {"service": "builder", "timeout_sec": 30, "retries": "three", "enabled": True},
        ensure_ascii=False,
        indent=2,
    ).encode("utf-8")


def _perf_xlsx() -> bytes:
    """`[XLSX-B]` — 직원명단.csv(`[CSV-A]`)와 `사번`으로 결합할 짝 파일."""
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "성과"
    sheet.append(["사번", "성과점수", "등급"])
    for row in [
        ["E001", 88, "A"],
        ["E002", 72, "B"],
        ["E003", 95, "A"],
        ["E004", 60, "C"],
    ]:
        sheet.append(row)
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _ops_guide_md() -> bytes:
    """`[MD-A]` — 제목 2개·번호 목록·굵은 글씨가 있는 UTF-8 Markdown."""
    text = (
        "# 운영 가이드\n\n"
        "## 배포 절차\n\n"
        "1. main 브랜치에서 태그를 만든다\n"
        "2. CI 통과를 확인한다\n"
        "3. **운영 승인** 후 배포를 실행한다\n\n"
        "## 롤백\n\n"
        "장애 시 직전 태그로 즉시 되돌린다. 자세한 내용은 담당자에게 문의한다.\n"
    )
    return text.encode("utf-8")


def _first_page_pdf(source_pdf: bytes) -> bytes:
    """`[PDF-WM]` — pdf_edit watermark 입력용 1페이지 PDF.

    LibreOffice(soffice)를 두 번 부르면 콜드 스타트가 겹쳐 타임아웃이 난다.
    이미 만든 PDF의 첫 페이지만 떼어 쓴다 — 워터마크 오버레이 입력은 페이지
    1장이면 충분하다.
    """
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(BytesIO(source_pdf))
    writer = PdfWriter()
    writer.add_page(reader.pages[0])
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _bundle_zip(members: dict[str, bytes]) -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in members.items():
            archive.writestr(name, data)
    return buffer.getvalue()


def _resolve_account(email: str | None) -> str:
    with database_connection() as connection, connection.cursor() as cursor:
        if email:
            cursor.execute(
                "SELECT account_id FROM user_account WHERE lower(email) = lower(%s)", (email,)
            )
        else:
            cursor.execute(
                "SELECT account_id FROM user_account WHERE account_status = 'ACTIVE' "
                "ORDER BY account_id LIMIT 1"
            )
        row = cursor.fetchone()
    if row is None:
        raise SystemExit(f"계정을 찾지 못했습니다: {email or '(ACTIVE 없음)'}")
    return row["account_id"]


def _purge_existing(account_id: str, names: list[str]) -> None:
    with database_connection() as connection, connection.cursor() as cursor:
        cursor.execute(
            """
            SELECT doc_id, storage_key FROM doc
            WHERE owner_account_id = %s AND source_type = 'UPLOAD'
              AND file_name = ANY(%s) AND deleted = false
            """,
            (account_id, names),
        )
        rows = list(cursor.fetchall())
    for row in rows:
        try:
            key = PersonalDocumentRepository.delete(
                doc_id=row["doc_id"], account_id=account_id
            )
            if key:
                storage.remove(key)
        except Exception as exc:  # noqa: BLE001
            print(f"  기존 {row['doc_id']} 정리 실패(무시): {exc}")


def _store(account_id: str, file_name: str, mime_type: str, data: bytes) -> str:
    doc_id = PersonalDocumentRepository.create(
        account_id=account_id, file_name=file_name, mime_type=mime_type
    )
    key = storage.build_personal_key(
        account_id=account_id, doc_id=doc_id, mime_type=mime_type
    )
    content_hash = storage.save(key, data)
    DocumentRepository.mark_stored(
        doc_id=doc_id,
        storage_key=key,
        content_hash=content_hash,
        revision=content_hash.removeprefix("sha256:")[:16],
    )
    if storage.is_download_only_upload(mime_type):
        PersonalDocumentRepository.set_search_enabled(
            doc_id=doc_id, account_id=account_id, enabled=False
        )
    return doc_id


def main() -> int:
    email = sys.argv[1] if len(sys.argv) > 1 else os.environ.get("SEED_ACCOUNT_EMAIL")
    account_id = _resolve_account(email)

    report = _report_docx(author="박작성 (원본 작성자)")
    report_revised = _report_docx(author="박작성 (원본 작성자)", revised=True)
    contract_pdf = office_to_pdf(data=report, mime_type=_DOCX)

    fixtures = [
        ("계약서_샘플.pdf", _PDF, contract_pdf),
        ("분기보고서.docx", _DOCX, report),
        ("분기보고서_수정본.docx", _DOCX, report_revised),
        ("부서별_공수.xlsx", _XLSX, _headcount_xlsx()),
        ("성과점수.xlsx", _XLSX, _perf_xlsx()),
        ("직원명단.csv", _CSV, _roster_csv()),
        ("설정.json", _JSON, _config_json()),
        ("운영가이드.md", _MD, _ops_guide_md()),
        ("워터마크.pdf", _PDF, _first_page_pdf(contract_pdf)),
        (
            "묶음.zip",
            _ZIP,
            _bundle_zip({"계약서_샘플.pdf": contract_pdf, "분기보고서.docx": report}),
        ),
    ]

    names = [name for name, _mime, _data in fixtures]
    _purge_existing(account_id, names)

    created = []
    for name, mime_type, data in fixtures:
        doc_id = _store(account_id, name, mime_type, data)
        created.append({"doc_id": doc_id, "file_name": name, "bytes": len(data)})
        print(f"  + {doc_id}  {name}  ({len(data)} bytes)")

    print(json.dumps({"account_id": account_id, "created": created}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

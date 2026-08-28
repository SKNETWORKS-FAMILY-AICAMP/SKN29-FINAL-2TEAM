"""Root 기본 Tool P0의 파일 처리 계약."""

from __future__ import annotations

import stat
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from django.test import SimpleTestCase
from docx import Document
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader, PdfWriter

from services.builtin_tools.common.errors import BuiltinToolError
from services.builtin_tools.data import compare_files
from services.builtin_tools.documents import (
    create_zip,
    edit_pdf,
    extract_zip,
    inspect_file,
    markdown_to_html,
    read_document,
    sanitize_file,
)
from services.builtin_tools.documents.reader import DOCX_MIME, PDF_MIME, XLSX_MIME


def _docx(*paragraphs: str, table: list[list[str]] | None = None, metadata: bool = False) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        target = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for column_index, value in enumerate(row):
                target.cell(row_index, column_index).text = value
    if metadata:
        document.core_properties.author = "작성자"
        document.core_properties.title = "비밀 제목"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx(values: dict[str, object], *, metadata: bool = False) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "데이터"
    for coordinate, value in values.items():
        sheet[coordinate] = value
    if metadata:
        workbook.properties.creator = "작성자"
        workbook.properties.title = "비밀 제목"
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _pdf(widths: list[float], *, metadata: bool = False, encrypted: bool = False) -> bytes:
    writer = PdfWriter()
    for width in widths:
        writer.add_blank_page(width=width, height=200)
    if metadata:
        writer.add_metadata({"/Author": "작성자", "/Title": "비밀 제목"})
    if encrypted:
        writer.encrypt("password")
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


class DocumentReadTests(SimpleTestCase):
    def test_reads_docx_paragraphs_and_tables(self):
        result = read_document(
            data=_docx("첫 문단", "둘째 문단", table=[["이름", "공수"], ["개발", "8"]]),
            mime_type=DOCX_MIME,
        )

        self.assertEqual([item["text"] for item in result["sections"]], ["첫 문단", "둘째 문단"])
        self.assertEqual(result["tables"][0]["rows"][1], ["개발", "8"])

    def test_reads_xlsx_values_and_formulas(self):
        result = read_document(data=_xlsx({"A1": "제목", "B2": "=1+1"}), mime_type=XLSX_MIME)

        rows = result["tables"][0]["rows"]
        self.assertEqual(rows[0][0], "제목")
        self.assertEqual(rows[1][1], "=1+1")

    def test_blank_pdf_returns_explicit_warning(self):
        result = read_document(data=_pdf([100]), mime_type=PDF_MIME)

        self.assertEqual(result["warnings"], [{"code": "TEXT_NOT_EXTRACTABLE", "page": 1}])

    def test_rejects_empty_and_unsupported_files(self):
        with self.assertRaisesRegex(BuiltinToolError, "빈 파일"):
            read_document(data=b"", mime_type=PDF_MIME)
        with self.assertRaisesRegex(BuiltinToolError, "PDF, DOCX, XLSX"):
            read_document(data=b"text", mime_type="text/plain")


class MarkdownConversionTests(SimpleTestCase):
    def test_converts_korean_markdown_without_executing_raw_html(self):
        result = markdown_to_html("# 제목\n\n- 항목\n\n<script>alert(1)</script>").decode()

        self.assertIn("<h1>제목</h1>", result)
        self.assertIn("<li>항목</li>", result)
        self.assertNotIn("<script>", result)
        self.assertIn("&lt;script&gt;", result)


class FileInspectAndSanitizeTests(SimpleTestCase):
    def test_inspects_hash_and_format_specific_counts(self):
        docx_result = inspect_file(data=_docx("본문", table=[["값"]]), mime_type=DOCX_MIME)
        xlsx_result = inspect_file(data=_xlsx({"A1": "값"}), mime_type=XLSX_MIME)
        pdf_result = inspect_file(data=_pdf([100, 200]), mime_type=PDF_MIME)

        self.assertEqual(len(docx_result["sha256"]), 64)
        self.assertEqual(docx_result["paragraph_count"], 1)
        self.assertEqual(docx_result["table_count"], 1)
        self.assertEqual(xlsx_result["sheet_count"], 1)
        self.assertEqual(pdf_result["page_count"], 2)
        self.assertTrue(pdf_result["mime_verified"])
        self.assertEqual(pdf_result["detected_mime_type"], PDF_MIME)

    def test_sanitizes_pdf_without_changing_pages(self):
        original = _pdf([100, 200], metadata=True)
        result = sanitize_file(data=original, mime_type=PDF_MIME)
        reader = PdfReader(BytesIO(result))

        self.assertEqual(len(reader.pages), 2)
        self.assertNotEqual(result, original)
        self.assertEqual(dict(reader.metadata or {}), {})

    def test_sanitizes_docx_without_changing_content(self):
        result = sanitize_file(data=_docx("보존할 본문", metadata=True), mime_type=DOCX_MIME)
        document = Document(BytesIO(result))

        self.assertEqual(document.paragraphs[0].text, "보존할 본문")
        self.assertEqual(document.core_properties.author, "")
        self.assertEqual(document.core_properties.title, "")

    def test_sanitizes_xlsx_without_changing_values_and_formulas(self):
        result = sanitize_file(
            data=_xlsx({"A1": "보존", "B1": "=1+1"}, metadata=True), mime_type=XLSX_MIME
        )
        workbook = load_workbook(BytesIO(result), data_only=False)

        self.assertEqual(workbook.active["A1"].value, "보존")
        self.assertEqual(workbook.active["B1"].value, "=1+1")
        self.assertIsNone(workbook.properties.creator)
        self.assertIsNone(workbook.properties.title)
        workbook.close()

    def test_rejects_encrypted_pdf(self):
        with self.assertRaisesRegex(BuiltinToolError, "암호"):
            sanitize_file(data=_pdf([100], encrypted=True), mime_type=PDF_MIME)

    def test_inspects_encrypted_pdf_without_decryption_error(self):
        result = inspect_file(data=_pdf([100], encrypted=True), mime_type=PDF_MIME)

        self.assertTrue(result["encrypted"])
        self.assertIsNone(result["page_count"])
        self.assertEqual(result["metadata"], {})

    def test_corrupt_ooxml_returns_stable_error(self):
        for mime_type in (DOCX_MIME, XLSX_MIME):
            with self.subTest(mime_type=mime_type):
                with self.assertRaises(BuiltinToolError):
                    read_document(data=b"PK\x03\x04broken", mime_type=mime_type)


class ArchiveTests(SimpleTestCase):
    def test_korean_file_names_round_trip(self):
        archive = create_zip([("보고서/업무.txt", b"task"), ("표.csv", b"a,b")])

        self.assertEqual(dict(extract_zip(archive)), {"보고서/업무.txt": b"task", "표.csv": b"a,b"})

    def test_rejects_traversal_and_duplicate_names(self):
        with self.assertRaisesRegex(BuiltinToolError, "안전하지 않은"):
            create_zip([("../secret.txt", b"x")])
        with self.assertRaisesRegex(BuiltinToolError, "중복"):
            create_zip([("same.txt", b"a"), ("same.txt", b"b")])
        with self.assertRaisesRegex(BuiltinToolError, "안전하지 않은"):
            create_zip([("safe.txt\x00hidden", b"x")])

    def test_rejects_symlink_entry(self):
        output = BytesIO()
        with ZipFile(output, "w", ZIP_DEFLATED) as archive:
            info = ZipInfo("link")
            info.create_system = 3
            info.external_attr = (stat.S_IFLNK | 0o777) << 16
            archive.writestr(info, "target")

        with self.assertRaisesRegex(BuiltinToolError, "심볼릭 링크"):
            extract_zip(output.getvalue())


class PdfEditTests(SimpleTestCase):
    def test_merge_and_reorder_preserve_requested_order(self):
        merged = edit_pdf(operation="merge", files=[_pdf([100]), _pdf([200])])
        reordered = edit_pdf(operation="reorder", files=[merged], pages=[2, 1])

        widths = [float(page.mediabox.width) for page in PdfReader(BytesIO(reordered)).pages]
        self.assertEqual(widths, [200.0, 100.0])

    def test_split_rotate_and_crop(self):
        original = _pdf([100, 200])
        split = edit_pdf(operation="split", files=[original], pages=[2])
        rotated = edit_pdf(operation="rotate", files=[original], pages=[1], rotation=90)
        cropped = edit_pdf(
            operation="crop", files=[original], pages=[1], crop_box=[0, 0, 50, 100]
        )

        self.assertEqual(len(split), 1)
        self.assertEqual(float(PdfReader(BytesIO(split[0])).pages[0].mediabox.width), 200.0)
        self.assertEqual(PdfReader(BytesIO(rotated)).pages[0].rotation, 90)
        self.assertEqual(float(PdfReader(BytesIO(cropped)).pages[0].cropbox.width), 50.0)

    def test_rejects_invalid_page_rotation_and_encrypted_pdf(self):
        source = _pdf([100])
        with self.assertRaisesRegex(BuiltinToolError, "범위"):
            edit_pdf(operation="extract", files=[source], pages=[2])
        with self.assertRaisesRegex(BuiltinToolError, "90, 180, 270"):
            edit_pdf(operation="rotate", files=[source], rotation=45)
        with self.assertRaisesRegex(BuiltinToolError, "암호"):
            edit_pdf(operation="merge", files=[_pdf([100], encrypted=True)])


class FileCompareTests(SimpleTestCase):
    def test_docx_detects_paragraph_and_table_changes(self):
        result = compare_files(
            before=_docx("기존", table=[["상태", "대기"]]),
            after=_docx("변경", table=[["상태", "완료"]]),
            mime_type=DOCX_MIME,
        )

        self.assertEqual(result["change_count"], 2)
        self.assertEqual({change["location"] for change in result["changes"]}, {"paragraphs", "tables"})

    def test_xlsx_reports_cell_locations_and_formula_changes(self):
        result = compare_files(
            before=_xlsx({"A1": "기존", "B1": "=1+1"}),
            after=_xlsx({"A1": "변경", "B1": "=2+2", "C1": "추가"}),
            mime_type=XLSX_MIME,
        )

        self.assertEqual(result["change_count"], 3)
        self.assertEqual({change["location"] for change in result["changes"]}, {"데이터!A1", "데이터!B1", "데이터!C1"})

    def test_identical_files_have_no_changes(self):
        data = _xlsx({"A1": "같음"})
        self.assertEqual(compare_files(before=data, after=data, mime_type=XLSX_MIME)["changes"], [])

    def test_xlsx_detects_added_and_removed_sheets(self):
        before_book = Workbook()
        before_book.active.title = "삭제할 시트"
        before_output = BytesIO()
        before_book.save(before_output)
        after_book = Workbook()
        after_book.active.title = "추가한 시트"
        after_output = BytesIO()
        after_book.save(after_output)

        result = compare_files(
            before=before_output.getvalue(),
            after=after_output.getvalue(),
            mime_type=XLSX_MIME,
        )

        self.assertEqual(
            {(change["type"], change["location"]) for change in result["changes"]},
            {("sheet_removed", "삭제할 시트"), ("sheet_added", "추가한 시트")},
        )

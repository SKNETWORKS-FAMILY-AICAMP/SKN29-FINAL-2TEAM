"""Root 기본 Tool P2의 시스템 문서 처리 계약."""

from __future__ import annotations

import hashlib
import signal
import subprocess
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from pathlib import Path
from unittest.mock import Mock, patch
from zipfile import ZIP_DEFLATED, ZipFile

import pdfplumber
from django.test import SimpleTestCase
from docx import Document
from pypdf import PdfReader

from services.builtin_tools.common.errors import BuiltinToolError
from services.builtin_tools.common.process import run_command
from services.builtin_tools.documents import (
    inspect_file,
    markdown_to_docx,
    markdown_to_pdf,
    office_to_pdf,
    sanitize_file,
)
from services.builtin_tools.documents.metadata import sensitive_metadata
from services.builtin_tools.documents.reader import DOCX_MIME, PDF_MIME, XLSX_MIME
from tests.test_builtin_tools_p0 import _docx, _pdf, _xlsx


def _pdf_text(data: bytes) -> str:
    with pdfplumber.open(BytesIO(data)) as document:
        return "\n".join(page.extract_text() or "" for page in document.pages)


def _with_relationship(data: bytes, relationship: bytes) -> bytes:
    source = BytesIO(data)
    output = BytesIO()
    with ZipFile(source) as original, ZipFile(output, "w", ZIP_DEFLATED) as changed:
        for info in original.infolist():
            content = original.read(info)
            if info.filename == "word/_rels/document.xml.rels":
                content = content.replace(b"</Relationships>", relationship + b"</Relationships>")
            changed.writestr(info, content)
    return output.getvalue()


def _with_high_ratio_entry(data: bytes) -> bytes:
    output = BytesIO()
    with ZipFile(BytesIO(data)) as original, ZipFile(output, "w", ZIP_DEFLATED) as changed:
        for info in original.infolist():
            changed.writestr(info, original.read(info))
        changed.writestr("word/media/compressed.bin", b"0" * (10 * 1024 * 1024))
    return output.getvalue()


class MimeDetectionTests(SimpleTestCase):
    def test_detects_pdf_docx_and_xlsx_from_bytes(self):
        fixtures = [
            (_pdf([100]), PDF_MIME),
            (_docx("본문"), DOCX_MIME),
            (_xlsx({"A1": "값"}), XLSX_MIME),
        ]

        for data, mime_type in fixtures:
            with self.subTest(mime_type=mime_type):
                result = inspect_file(data=data, mime_type=mime_type)
                self.assertEqual(result["detected_mime_type"], mime_type)
                self.assertTrue(result["mime_verified"])

    def test_reports_declared_and_actual_mime_mismatch(self):
        result = inspect_file(data=_pdf([100]), mime_type=DOCX_MIME)

        self.assertEqual(result["declared_mime_type"], DOCX_MIME)
        self.assertEqual(result["detected_mime_type"], PDF_MIME)
        self.assertFalse(result["mime_verified"])
        self.assertEqual(result["page_count"], 1)


class SystemConversionTests(SimpleTestCase):
    docx_data = _docx(
        "한국어 업무 보고서",
        "완료 기준을 확인합니다.",
        table=[["항목", "값"], ["일정", "완료"]],
    )
    xlsx_data = _xlsx({"A1": "한국어 업무표", "A2": "개발", "B2": 8})
    markdown = "# 한국어 제목\n\n- 첫 번째 항목\n\n| 항목 | 값 |\n| --- | --- |\n| 일정 | 완료 |"

    def test_docx_and_xlsx_convert_to_readable_korean_pdf(self):
        docx_pdf = office_to_pdf(data=self.docx_data, mime_type=DOCX_MIME)
        xlsx_pdf = office_to_pdf(data=self.xlsx_data, mime_type=XLSX_MIME)

        self.assertIn("한국어 업무 보고서", _pdf_text(docx_pdf))
        self.assertIn("완료 기준", _pdf_text(docx_pdf))
        self.assertIn("한국어 업무표", _pdf_text(xlsx_pdf))
        self.assertGreaterEqual(len(PdfReader(BytesIO(docx_pdf)).pages), 1)

    def test_markdown_converts_to_docx_and_pdf_with_table(self):
        docx = markdown_to_docx(self.markdown)
        pdf = markdown_to_pdf(self.markdown)
        document = Document(BytesIO(docx))

        self.assertEqual(document.paragraphs[0].text, "한국어 제목")
        self.assertEqual(document.tables[0].cell(1, 0).text, "일정")
        self.assertIn("한국어 제목", _pdf_text(pdf))
        self.assertIn("완료", _pdf_text(pdf))

    def test_rejects_images_external_office_resources_and_mime_spoofing(self):
        with self.assertRaisesRegex(BuiltinToolError, "이미지"):
            markdown_to_docx("![외부 이미지](https://example.invalid/image.png)")
        with self.assertRaisesRegex(BuiltinToolError, "외부 파일"):
            office_to_pdf(
                data=_with_relationship(
                    self.docx_data,
                    b'<Relationship Id="external-image" '
                    b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
                    b'Target="https://example.invalid/image.png" TargetMode="External"/>',
                ),
                mime_type=DOCX_MIME,
            )
        with self.assertRaisesRegex(BuiltinToolError, "문서 밖"):
            office_to_pdf(
                data=_with_relationship(
                    self.docx_data,
                    b'<Relationship Id="unsafe-file" '
                    b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
                    b'Target="../../../../etc/passwd"/>',
                ),
                mime_type=DOCX_MIME,
            )
        with self.assertRaisesRegex(BuiltinToolError, "압축 크기"):
            office_to_pdf(data=_with_high_ratio_entry(self.docx_data), mime_type=DOCX_MIME)
        with self.assertRaisesRegex(BuiltinToolError, "실제 파일 형식"):
            office_to_pdf(data=_pdf([100]), mime_type=DOCX_MIME)

    def test_original_is_unchanged_and_four_conversions_do_not_collide(self):
        original_hash = hashlib.sha256(self.docx_data).hexdigest()
        with ThreadPoolExecutor(max_workers=4) as executor:
            outputs = list(
                executor.map(
                    lambda _: office_to_pdf(data=self.docx_data, mime_type=DOCX_MIME),
                    range(4),
                )
            )

        self.assertEqual(hashlib.sha256(self.docx_data).hexdigest(), original_hash)
        self.assertTrue(all(_pdf_text(output) == _pdf_text(outputs[0]) for output in outputs))

    def test_missing_system_dependency_returns_stable_error(self):
        with patch("services.builtin_tools.common.process.shutil.which", return_value=None):
            with self.assertRaisesRegex(BuiltinToolError, "설치"):
                markdown_to_docx("# 제목")

    def test_process_timeout_kills_the_process_group(self):
        process = Mock(pid=12345, returncode=None)
        process.communicate.side_effect = [
            subprocess.TimeoutExpired(cmd="pandoc", timeout=1),
            (b"", b""),
        ]
        with (
            patch("services.builtin_tools.common.process.shutil.which", return_value="/usr/bin/pandoc"),
            patch("services.builtin_tools.common.process.subprocess.Popen", return_value=process),
            patch("services.builtin_tools.common.process.os.killpg") as kill_group,
        ):
            with self.assertRaisesRegex(BuiltinToolError, "제한 시간"):
                run_command(
                    "pandoc",
                    [],
                    cwd=Path("/tmp"),
                    home=Path("/tmp"),
                    timeout_seconds=1,
                )

        kill_group.assert_called_once_with(12345, signal.SIGKILL)


class ExifMetadataTests(SimpleTestCase):
    def test_sensitive_metadata_is_removed_from_all_supported_formats(self):
        fixtures = [
            (_pdf([100], metadata=True), PDF_MIME),
            (_docx("본문", metadata=True), DOCX_MIME),
            (_xlsx({"A1": "값"}, metadata=True), XLSX_MIME),
        ]

        for original, mime_type in fixtures:
            with self.subTest(mime_type=mime_type):
                self.assertTrue(sensitive_metadata(data=original, mime_type=mime_type))
                sanitized = sanitize_file(data=original, mime_type=mime_type)
                self.assertEqual(sensitive_metadata(data=sanitized, mime_type=mime_type), {})

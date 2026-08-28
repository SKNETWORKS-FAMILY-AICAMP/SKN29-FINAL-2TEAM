"""Root 기본 Tool P1의 표·계산 처리 계약."""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

from django.test import SimpleTestCase
from docx import Document
from openpyxl import Workbook, load_workbook

from services.builtin_tools.calculation import calculate
from services.builtin_tools.common.errors import BuiltinToolError
from services.builtin_tools.data import check_data_quality, compare_files, transform_table
from services.builtin_tools.data.transformer import CSV_MIME, JSON_MIME, PARQUET_MIME
from services.builtin_tools.documents import document_to_markdown, read_document
from services.builtin_tools.documents.reader import DOCX_MIME, PDF_MIME, XLSX_MIME


def _csv(text: str) -> bytes:
    return text.strip().encode("utf-8")


def _docx() -> bytes:
    document = Document()
    document.add_paragraph("업무 현황")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "담당"
    table.cell(0, 1).text = "공수"
    table.cell(1, 0).text = "개발"
    table.cell(1, 1).text = "8"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "업무"
    sheet.append(["담당", "공수"])
    sheet.append(["개발", 8])
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


class DocumentP1Tests(SimpleTestCase):
    def test_real_korean_pdf_text_and_tables_are_extracted(self):
        path = (
            Path(__file__).parent
            / "eval/documents/pdf/한빛몰_주문정산_요구사항정의서.pdf"
        )
        result = read_document(data=path.read_bytes(), mime_type=PDF_MIME)

        text = " ".join(section["text"] for section in result["sections"])
        self.assertIn("요구사항", text)
        self.assertGreaterEqual(len(result["tables"]), 1)

    def test_document_to_markdown_supports_docx_and_xlsx(self):
        docx_markdown = document_to_markdown(data=_docx(), mime_type=DOCX_MIME)
        xlsx_markdown = document_to_markdown(data=_xlsx(), mime_type=XLSX_MIME)

        self.assertIn("업무 현황", docx_markdown)
        self.assertIn("| 담당 | 공수 |", docx_markdown)
        self.assertIn("## 업무", xlsx_markdown)

    def test_pdf_comparison_detects_text_change(self):
        base = Path(__file__).parent / "eval/documents/pdf"
        result = compare_files(
            before=(base / "한빛몰_주문정산_요구사항정의서.pdf").read_bytes(),
            after=(base / "한빛몰_주문정산_기술검토회의록.pdf").read_bytes(),
            mime_type=PDF_MIME,
        )

        self.assertGreater(result["change_count"], 0)
        self.assertTrue(any(change["location"] == "pages" for change in result["changes"]))


class TableTransformTests(SimpleTestCase):
    data = _csv("팀,담당,공수\nA,가,3\nA,나,7\nB,다,5")

    def test_filters_sorts_and_selects_without_sql_input(self):
        result = transform_table(
            data=self.data,
            mime_type=CSV_MIME,
            operation="filter",
            columns=["담당", "공수"],
            filters=[{"column": "공수", "operator": "gte", "value": 5}],
            sort=[{"column": "공수", "direction": "desc"}],
        )

        self.assertEqual(result["rows"], [["나", 7], ["다", 5]])

        literal_wildcard = transform_table(
            data=_csv("name\na%b\naxb"),
            mime_type=CSV_MIME,
            operation="filter",
            filters=[{"column": "name", "operator": "contains", "value": "%"}],
        )
        self.assertEqual(literal_wildcard["rows"], [["a%b"]])

    def test_aggregates_and_joins(self):
        aggregate = transform_table(
            data=self.data,
            mime_type=CSV_MIME,
            operation="aggregate",
            group_by=["팀"],
            metrics=[{"function": "sum", "column": "공수", "as": "합계"}],
            sort=[{"column": "팀", "direction": "asc"}],
        )
        joined = transform_table(
            data=self.data,
            mime_type=CSV_MIME,
            operation="join",
            join_data=_csv("담당,직무\n가,개발\n다,기획"),
            join_mime_type=CSV_MIME,
            join_on=[{"left": "담당", "right": "담당"}],
        )

        self.assertEqual(aggregate["rows"], [["A", 10], ["B", 5]])
        self.assertEqual(joined["row_count"], 2)
        self.assertIn("right.직무", joined["columns"])

    def test_tolerates_metrics_and_group_by_shape_mistakes(self):
        """모델이 배열 대신 객체 하나나 문자열을 보내도 배열로 맞춘다(2026-08-28).

        전에는 `metrics={"공수": {...}}` 를 보내면 `TOOL_EXECUTION_FAILED`(내부
        예외)로 죽었다.
        """

        by_key = transform_table(
            data=self.data,
            mime_type=CSV_MIME,
            operation="aggregate",
            group_by="팀",
            metrics={"공수": {"function": "sum", "alias": "합계"}},
        )
        by_dict = transform_table(
            data=self.data,
            mime_type=CSV_MIME,
            operation="aggregate",
            group_by=["팀"],
            metrics={"column": "공수", "function": "sum", "alias": "합계"},
        )
        self.assertEqual(by_key["rows"], [["A", 10], ["B", 5]])
        self.assertEqual(by_dict["rows"], [["A", 10], ["B", 5]])

        with self.assertRaisesRegex(BuiltinToolError, "객체 배열"):
            transform_table(
                data=self.data,
                mime_type=CSV_MIME,
                operation="aggregate",
                group_by=["팀"],
                metrics="공수",
            )

    def test_exports_json_xlsx_and_parquet(self):
        json_result = transform_table(
            data=self.data, mime_type=CSV_MIME, operation="convert", output_format="json"
        )
        xlsx_result = transform_table(
            data=self.data, mime_type=CSV_MIME, operation="convert", output_format="xlsx"
        )
        parquet_result = transform_table(
            data=self.data, mime_type=CSV_MIME, operation="convert", output_format="parquet"
        )

        self.assertEqual(len(json.loads(json_result["output_bytes"])), 3)
        workbook = load_workbook(BytesIO(xlsx_result["output_bytes"]), read_only=True)
        self.assertEqual(workbook.active.max_row, 4)
        workbook.close()
        reopened = transform_table(
            data=parquet_result["output_bytes"],
            mime_type=PARQUET_MIME,
            operation="preview",
        )
        self.assertEqual(reopened["row_count"], 3)

    def test_export_is_not_cut_to_preview_limit_and_statistics_work(self):
        exported = transform_table(
            data=self.data,
            mime_type=CSV_MIME,
            operation="convert",
            limit=1,
            output_format="json",
        )
        statistics = transform_table(
            data=self.data, mime_type=CSV_MIME, operation="statistics"
        )

        self.assertTrue(exported["truncated"])
        self.assertEqual(len(json.loads(exported["output_bytes"])), 3)
        self.assertEqual(statistics["row_count"], 3)

    def test_csv_json_and_parquet_aggregation_results_match(self):
        payload = [
            {"팀": "A", "공수": 3},
            {"팀": "A", "공수": 7},
            {"팀": "B", "공수": 5},
        ]
        parquet = transform_table(
            data=self.data,
            mime_type=CSV_MIME,
            operation="convert",
            output_format="parquet",
        )["output_bytes"]
        sources = [
            (self.data, CSV_MIME),
            (json.dumps(payload, ensure_ascii=False).encode(), JSON_MIME),
            (parquet, PARQUET_MIME),
        ]
        results = []
        for data, mime_type in sources:
            result = transform_table(
                data=data,
                mime_type=mime_type,
                operation="aggregate",
                group_by=["팀"],
                metrics=[{"function": "sum", "column": "공수", "as": "합계"}],
                sort=[{"column": "팀", "direction": "asc"}],
            )
            results.append(result["rows"])

        self.assertEqual(results, [[["A", 10], ["B", 5]]] * 3)

    def test_rejects_unknown_columns_and_unsupported_operations(self):
        with self.assertRaisesRegex(BuiltinToolError, "없는 열"):
            transform_table(
                data=self.data, mime_type=CSV_MIME, operation="preview", columns=["DROP TABLE"]
            )
        with self.assertRaisesRegex(BuiltinToolError, "지원하지 않는 표 가공"):
            transform_table(data=self.data, mime_type=CSV_MIME, operation="raw_sql")

    def test_reads_xlsx_and_json_sources(self):
        xlsx_result = transform_table(data=_xlsx(), mime_type=XLSX_MIME, operation="preview")
        json_result = transform_table(
            data=json.dumps([{"담당": "개발", "공수": 8}], ensure_ascii=False).encode(),
            mime_type=JSON_MIME,
            operation="preview",
        )

        self.assertEqual(xlsx_result["rows"], [["개발", 8]])
        self.assertEqual(json_result["rows"], [["개발", 8]])


class DataQualityTests(SimpleTestCase):
    def test_reports_missing_duplicate_type_and_range_errors(self):
        result = check_data_quality(
            data=_csv("id,score\n1,90\n1,90\n2,\n3,wrong"),
            mime_type=CSV_MIME,
            schema={
                "fields": [
                    {"name": "id", "type": "integer", "constraints": {"required": True}},
                    {
                        "name": "score",
                        "type": "integer",
                        "constraints": {"required": True, "minimum": 0, "maximum": 100},
                    },
                ]
            },
        )

        codes = {error["code"] for error in result["errors"]}
        self.assertFalse(result["valid"])
        self.assertIn("missing-value", codes)
        self.assertIn("duplicate-row", codes)
        self.assertTrue(codes & {"type-error", "constraint-error", "required-constraint"})

    def test_valid_data_infers_schema_without_false_errors(self):
        result = check_data_quality(
            data=_csv("id,name\n1,가\n2,나"), mime_type=CSV_MIME
        )

        self.assertTrue(result["valid"])
        self.assertEqual(result["error_count"], 0)
        self.assertEqual([field["name"] for field in result["inferred_schema"]["fields"]], ["id", "name"])

    def test_json_schema_errors_include_location(self):
        result = check_data_quality(
            data=json.dumps([{"name": 3}]).encode(),
            mime_type=JSON_MIME,
            json_schema={
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {"name": {"type": "string"}},
                    "required": ["name"],
                },
            },
        )

        error = next(item for item in result["errors"] if item["code"] == "json-schema")
        self.assertEqual(error["location"], [0, "name"])

    def test_rejects_empty_and_non_utf8_csv(self):
        with self.assertRaisesRegex(BuiltinToolError, "빈 표"):
            check_data_quality(data=b"", mime_type=CSV_MIME)
        with self.assertRaisesRegex(BuiltinToolError, "UTF-8"):
            check_data_quality(data=b"\xff\xfe", mime_type=CSV_MIME)

    def test_single_json_object_is_validated_against_json_schema(self):
        """설정 파일 한 건을 스키마에 대조하는 것은 표가 아니어도 된다(2026-08-28)."""

        config = json.dumps({"retries": "three", "timeout_sec": 30}).encode()
        schema = {
            "type": "object",
            "properties": {"retries": {"type": "integer"}},
            "required": ["retries"],
        }

        result = check_data_quality(data=config, mime_type=JSON_MIME, json_schema=schema)

        self.assertFalse(result["valid"])
        self.assertEqual(result["errors"][0]["code"], "json-schema")
        self.assertEqual(result["errors"][0]["location"], ["retries"])

        # 스키마 없이 표가 아닌 JSON 은 여전히 거절한다 — 그때는 검사할 기준이 없다.
        with self.assertRaisesRegex(BuiltinToolError, "객체 배열"):
            check_data_quality(data=config, mime_type=JSON_MIME)


class CalculatorTests(SimpleTestCase):
    def test_math_is_exact_and_rejects_code_execution(self):
        result = calculate(operation="math", expression="sqrt(16) + 1 / 3")
        percentage = calculate(operation="math", expression="500 * percent(20)")

        self.assertEqual(result["exact"], "13/3")
        self.assertEqual(percentage["decimal"], 100)
        with self.assertRaisesRegex(BuiltinToolError, "허용된"):
            calculate(operation="math", expression="__import__('os').system('id')")
        with self.assertRaisesRegex(BuiltinToolError, "지수"):
            calculate(operation="math", expression="2 ** 1000")
        with self.assertRaisesRegex(BuiltinToolError, "실수"):
            calculate(operation="math", expression="sqrt(-1)")

    def test_converts_units_including_offset_temperature(self):
        distance = calculate(
            operation="unit", amount=1, from_unit="kilometer", to_unit="meter"
        )
        temperature = calculate(
            operation="unit", amount=0, from_unit="degC", to_unit="degF"
        )

        self.assertEqual(distance["value"], 1000)
        self.assertAlmostEqual(temperature["value"], 32)

    def test_results_are_tidied_of_float_noise(self):
        # 25°C → 77°F 는 부동소수로 76.99999999999993 이 나온다(2026-08-29 정리).
        self.assertEqual(
            calculate(operation="unit", amount=25, from_unit="degC", to_unit="degF")["value"],
            77.0,
        )
        # 의미 있는 소수는 유지한다.
        mile = calculate(operation="unit", amount=10, from_unit="km", to_unit="mile")["value"]
        self.assertAlmostEqual(mile, 6.2137119, places=6)
        self.assertEqual(calculate(operation="math", expression="(3+4)*5 - 2**3")["decimal"], 27.0)

    def test_parses_korean_dates_deterministically(self):
        leap_day = calculate(
            operation="date", expression="2028년 2월 29일", relative_base="2026-08-27"
        )
        next_friday = calculate(
            operation="date", expression="다음 주 금요일", relative_base="2026-08-27"
        )

        self.assertTrue(leap_day["datetime"].startswith("2028-02-29"))
        self.assertTrue(next_friday["datetime"].startswith("2026-09-04"))

    def test_duration_business_days_and_timezone(self):
        duration = calculate(
            operation="duration", start="2026-08-27", end="2026-08-29"
        )
        business = calculate(
            operation="business_days",
            start_date="2026-08-24",
            end_date="2026-08-28",
            country="KR",
            company_holidays=["2026-08-26"],
        )
        timezone = calculate(
            operation="timezone",
            datetime="2026-08-27T09:00:00",
            from_timezone="Asia/Seoul",
            to_timezone="UTC",
        )

        self.assertEqual(duration["days"], 2)
        self.assertEqual(business["business_days"], 4)
        self.assertTrue(timezone["datetime"].startswith("2026-08-27T00:00:00+00:00"))

        dst = calculate(
            operation="timezone",
            datetime="2026-03-08T07:30:00+00:00",
            from_timezone="UTC",
            to_timezone="America/New_York",
        )
        self.assertTrue(dst["datetime"].startswith("2026-03-08T03:30:00-04:00"))

        with self.assertRaisesRegex(BuiltinToolError, "시간대 표기"):
            calculate(
                operation="duration",
                start="2026-08-27T09:00:00+09:00",
                end="2026-08-27T10:00:00",
            )

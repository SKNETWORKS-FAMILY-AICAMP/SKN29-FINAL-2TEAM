"""`document_create` 가 굽는 docx 의 계약.

여기서 보는 것은 **모델이 쓴 마크다운 중 무엇을 그리고 무엇을 안 그리는가**다.
해석 범위를 좁게 고정하되, 검증 가능한 HTTP(S) 출처는 표시 이름을 보존한
클릭 가능한 링크로 만든다.
"""

from io import BytesIO
from zipfile import ZipFile
from unittest.mock import patch

from django.test import SimpleTestCase
from docx import Document
from docx.oxml.ns import qn

from services.document_export import build_docx
from services.harness.registry import _DOCX_MIME


def _doc(data: bytes):
    return Document(BytesIO(data))


def _paras(data: bytes):
    return [(p.style.name, p.text) for p in _doc(data).paragraphs if p.text.strip()]


class BuildDocxTests(SimpleTestCase):
    def test_제목이_맨_위에_굵게_들어간다(self):
        paras = _paras(build_docx(title="주간 보고", body="본문"))

        self.assertEqual(paras[0][1], "주간 보고")
        self.assertTrue(_doc(build_docx(title="주간 보고", body="본문")).paragraphs[0].runs[0].bold)

    def test_제목_수준_셋을_구분한다(self):
        paras = _paras(build_docx(title="t", body="# 하나\n## 둘\n### 셋"))

        self.assertEqual(
            [(s, t) for s, t in paras[1:]],
            [("Heading 1", "하나"), ("Heading 2", "둘"), ("Heading 3", "셋")],
        )

    def test_목록_두_종류를_구분한다(self):
        paras = _paras(build_docx(title="t", body="- 글머리\n* 같은것\n1. 번호\n2) 도번호"))

        self.assertEqual(
            [s for s, _ in paras[1:]],
            ["List Bullet", "List Bullet", "List Number", "List Number"],
        )
        self.assertEqual([t for _, t in paras[1:]], ["글머리", "같은것", "번호", "도번호"])

    def test_굵게와_http_링크를_살리고_나머지_문법은_글자_그대로_둔다(self):
        data = build_docx(title="t", body="**핵심**은 배포다. `코드` 와 [공식 문서](https://example.com/docs)를 본다.")
        body = _doc(data).paragraphs[1]

        self.assertEqual([(r.text, r.bold) for r in body.runs][0], ("핵심", True))
        # 백틱은 그대로 두되, 링크 문법과 원시 URL은 본문에서 숨긴다.
        self.assertIn("`코드`", body.text)
        with ZipFile(BytesIO(data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
            relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        self.assertIn("공식 문서", document_xml)
        self.assertNotIn("[공식 문서]", document_xml)
        self.assertIn("https://example.com/docs", relationships)

    def test_구조화_표의_출처_이름도_클릭_가능하게_만든다(self):
        data = build_docx(
            title="출처 표",
            blocks=[
                {
                    "type": "table",
                    "headers": ["항목", "출처"],
                    "rows": [["Claude", "[Anthropic 공식 홈페이지](https://www.anthropic.com)"]],
                }
            ],
        )

        with ZipFile(BytesIO(data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
            relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        self.assertIn("Anthropic 공식 홈페이지", document_xml)
        self.assertIn("https://www.anthropic.com", relationships)

    def test_빈_줄은_빈_문단을_만들지_않는다(self):
        """워드는 문단 간격을 이미 주므로 빈 문단을 넣으면 여백이 두 배가 된다."""

        doc = _doc(build_docx(title="t", body="첫 문단\n\n\n\n둘째 문단"))

        self.assertEqual([p.text for p in doc.paragraphs], ["t", "첫 문단", "둘째 문단"])

    def test_본문이_비어도_제목만_있는_파일이_나온다(self):
        self.assertEqual([p.text for p in _doc(build_docx(title="t", body="")).paragraphs], ["t"])

    def test_제목이_비면_지어내지_않고_자리만_채운다(self):
        self.assertEqual(_paras(build_docx(title="   ", body="본문"))[0][1], "제목 없음")

    def test_제어문자가_섞여도_열리는_파일이_나온다(self):
        """워드가 XML 에 못 담는 값이 본문에 있으면 파일 자체가 안 열린다."""

        data = build_docx(title="t\x00", body="앞\x0b뒤")

        self.assertEqual([p.text for p in _doc(data).paragraphs], ["t", "앞뒤"])

    def test_회사_문서용_기본_서식과_페이지_정보를_넣는다(self):
        data = build_docx(title="주간 보고", body="# 요약\n본문")
        doc = _doc(data)

        self.assertEqual(doc.core_properties.title, "주간 보고")
        self.assertEqual(doc.sections[0].header.paragraphs[0].text, "주간 보고")
        self.assertIn("페이지", doc.sections[0].footer.paragraphs[0].text)
        self.assertEqual(doc.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")), "맑은 고딕")
        self.assertEqual(doc.styles["Heading 1"].font.color.rgb, doc.paragraphs[0].runs[0].font.color.rgb)
        self.assertTrue(doc.styles["Heading 1"].paragraph_format.keep_with_next)
        with ZipFile(BytesIO(data)) as archive:
            footer_xml = archive.read("word/footer1.xml").decode("utf-8")
            settings_xml = archive.read("word/settings.xml").decode("utf-8")
        self.assertIn(" PAGE ", footer_xml)
        self.assertNotIn("updateFields", settings_xml)

    def test_업무보고서_메타데이터와_구조화_블록을_그린다(self):
        data = build_docx(
            title="주간 업무 보고",
            template_id="business_report",
            metadata={"department": "개발팀", "author": "임준억", "date": "2026-08-28"},
            blocks=[
                {"type": "heading", "level": 1, "text": "핵심 요약"},
                {"type": "note", "label": "결론", "text": "승인 전에 내용을 확인합니다."},
                {"type": "bullet_list", "items": ["문서 생성", "미리보기 검증"]},
                {
                    "type": "table",
                    "headers": ["항목", "상태"],
                    "rows": [["문서 생성", "완료"], ["배포", "대기"]],
                },
            ],
        )
        doc = _doc(data)

        self.assertEqual(doc.tables[0].cell(0, 0).text, "부서")
        self.assertEqual(doc.tables[0].cell(0, 1).text, "개발팀")
        self.assertEqual(doc.tables[1].cell(0, 0).text, "항목")
        self.assertEqual(doc.tables[1].cell(2, 1).text, "대기")
        self.assertIn(("Heading 1", "핵심 요약"), _paras(data))
        self.assertIn(("List Bullet", "문서 생성"), _paras(data))

    def test_구조화_표의_잘못된_행은_거절한다(self):
        with self.assertRaisesRegex(ValueError, "열 개수"):
            build_docx(
                title="잘못된 표",
                blocks=[{"type": "table", "headers": ["하나"], "rows": [[1, 2]]}],
            )


class DocumentCreateToolTests(SimpleTestCase):
    def _call(self, **kwargs):
        from services.harness.registry import _document_create

        defaults = {"account_id": "AC001", "title": "주간 보고", "body": "# 요약\n본문"}
        return _document_create(**{**defaults, **kwargs})

    @patch("services.harness.registry.DocumentRepository.mark_stored")
    @patch("services.harness.registry.storage.save", return_value="sha256:abc123def4567890")
    @patch(
        "services.harness.registry.PersonalDocumentRepository.create_generated",
        return_value="DC100",
    )
    def test_생성_문서로_저장한다(self, create, save, mark_stored):
        result = self._call()

        # `table_export` 와 같은 자리·같은 규칙이다.
        self.assertEqual(create.call_args.kwargs["mime_type"], _DOCX_MIME)
        self.assertTrue(create.call_args.kwargs["file_name"].endswith(".docx"))
        self.assertTrue(save.call_args.args[0].startswith("user/AC001/DC100"))
        self.assertTrue(save.call_args.args[0].endswith(".docx"))

        self.assertEqual([p.text for p in _doc(save.call_args.args[1]).paragraphs][:2], ["주간 보고", "요약"])
        # `table_export` 와 같은 `file` 계약을 쓴다.
        self.assertEqual(
            result["file"], {"doc_id": "DC100", "file_name": create.call_args.kwargs["file_name"], "mime_type": _DOCX_MIME}
        )
        # 본문을 되돌려주지 않는다 — 모델이 방금 보낸 값이라 컨텍스트만 두 배가 된다.
        self.assertNotIn("body", result)

    def test_사람이_고칠_수_있는_실패는_그대로_말한다(self):
        from services.harness.registry import ToolInputError

        with self.assertRaises(ToolInputError):
            self._call(body="")
        with self.assertRaises(ToolInputError):
            self._call(body="   \n\n  ")

    @patch("services.harness.registry.DocumentRepository.mark_stored")
    @patch("services.harness.registry.storage.save", return_value="sha256:abc123def4567890")
    @patch(
        "services.harness.registry.PersonalDocumentRepository.create_generated",
        return_value="DC101",
    )
    def test_구조화_블록만으로도_문서를_만든다(self, create, save, mark_stored):
        result = self._call(
            body="",
            template_id="business_report",
            blocks=[{"type": "paragraph", "text": "구조화 본문"}],
        )

        self.assertEqual(result["file"]["doc_id"], "DC101")
        self.assertIn("구조화 본문", [p.text for p in _doc(save.call_args.args[1]).paragraphs])
